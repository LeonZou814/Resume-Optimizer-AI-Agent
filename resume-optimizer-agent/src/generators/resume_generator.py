"""
简历生成模块
将优化后的内容生成专业的 DOCX / Markdown 格式
"""

import re
from pathlib import Path
from typing import Dict, Optional, List
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from config.config import settings


# ── LLM 输出清洗 ──────────────────────────────────────────

# 常见的 LLM 元文本前缀（需要去除的行）
_META_PATTERNS = [
    r"^润色后的.*[：:]",
    r"^优化后的.*[：:]",
    r"^以下是.*[：:]",
    r"^严格基于原文.*",
    r"^（严格基于.*）",
    r"^注意[：:].*不可超出",
    r"^说明[：:].*事实来源",
    # LLM 末尾的自证说明（括号包裹的整行注释）
    r"^（注[：:].*）$",
    r"^（说明[：:].*）$",
    r"^（以上内容均基于.*）$",
    r"^（基于原始简历.*）$",
    r"^（基于原文.*）$",
    # 不带括号的注释
    r"^注[：:].*未添加.*",
    r"^注[：:].*源自.*",
    r"^注[：:].*事实来源.*",
    r"^【注[：:].*】$",
    r"^【说明[：:].*】$",
    # LLM 末尾的自证说明（含关键词组合）
    r"^原文未提及.*避免.*虚构.*",
    r"^严格依据事实.*",
    r"^【建议学习】.*",
    # "以上" 开头的总结性注释
    r"^以上优化.*原文.*",
    r"^以上.*均源自.*",
    r"^以上.*均基于.*原文.*",
    # "本次" 开头的自证说明
    r"^本次优化仅基于.*",
    r"^本次.*未添加.*",
    # 含 "未添加" + "虚构" 的任意行
    r"^.*未添加.*虚构.*$",
    r"^.*所有数据均源自.*",
    r"^.*所有内容均.*原文.*$",
    # LLM 输出的"重要提醒"、"注意事项"等元文本
    r"^重要提醒[：:].*",
    r"^注意事项[：:].*",
    r"^建议修改[：:].*",
    r"^温馨提示[：:].*",
    r"^⚠.*",
    r"^📌.*",
    # 多行元文本块（以特定关键词开头的整段）
    r"^【重要提醒】.*",
    r"^【注意事项】.*",
    r"^【建议】.*",
]


# 触发"后续全部跳过"的元文本模式（一旦匹配，后续所有行都视为注释）
_META_BLOCK_PATTERNS = [
    r"^重要提醒[：:]",
    r"^⚠.*",
    r"^📌.*",
    r"^【重要提醒】",
    r"^【注意事项】",
    r"^【建议】",
]


def _clean_llm_output(text: str) -> str:
    """
    清洗 LLM 输出，去除元文本前缀和说明性文字，只保留简历实际内容。
    """
    if not text:
        return text

    lines = text.split('\n')
    cleaned = []
    in_meta_block = False  # 是否进入了元文本块（后续全部跳过）

    for line in lines:
        stripped = line.strip()

        # 一旦检测到元文本块起始行，后续所有行全部跳过
        if in_meta_block:
            continue

        # 检查是否进入元文本块
        for pattern in _META_BLOCK_PATTERNS:
            if re.match(pattern, stripped):
                in_meta_block = True
                break
        if in_meta_block:
            continue

        # 跳过空行（连续多个空行只保留一个）
        if not stripped:
            if cleaned and cleaned[-1] != "":
                cleaned.append("")
            continue

        # 检查是否是单行元文本前缀
        is_meta = False
        for pattern in _META_PATTERNS:
            if re.match(pattern, stripped):
                is_meta = True
                break
        if is_meta:
            continue

        cleaned.append(line)

    # 去除首尾空行
    result = '\n'.join(cleaned).strip()

    # 修复不完整的加粗标记 *text** → **text**（LLM 有时会漏掉一个前导星号）
    # 仅匹配恰好1个前导*（前面不是*，第3个字符也不是*）+ 文本 + **结尾
    result = re.sub(
        r'(?<!\*)\*(?!\*)([^\s*]+?\*\*)(?!\*)',
        lambda m: m.group(0) if m.group(1).startswith('**') else '**' + m.group(1),
        result
    )

    return result


# ── 简历生成器 ──────────────────────────────────────────

class ResumeGenerator:
    """简历生成器"""

    # 配色方案
    PRIMARY_COLOR = RGBColor(0x1A, 0x56, 0x8E)     # 深蓝 - 标题
    SECONDARY_COLOR = RGBColor(0x33, 0x33, 0x33)    # 深灰 - 正文
    ACCENT_COLOR = RGBColor(0x2E, 0x86, 0xC1)       # 亮蓝 - 强调
    LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)          # 浅灰 - 辅助信息
    SECTION_BG = "1A568E"                             # 标题栏背景色 (hex)

    def __init__(self, template_dir: str = None):
        self.template_dir = Path(template_dir) if template_dir else Path("data/templates")
        self.output_dir = Path(settings.output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_docx(self, optimized_data: Dict, original_resume: Dict, filename: str = None) -> str:
        """生成专业排版的 DOCX 简历"""
        if filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"optimized_resume_{timestamp}.docx"

        output_path = self.output_dir / filename
        doc = Document()

        # ── 页面设置 ──
        self._setup_page(doc)

        # ── 默认样式 ──
        self._setup_styles(doc)

        # ── 姓名（大标题） ──
        name = original_resume.get("name") or self._extract_name(original_resume)
        if name:
            self._add_name_header(doc, name)

        # ── 联系方式（一行居中） ──
        self._add_contact_row(doc, original_resume)

        # ── 各模块（优先保持原始简历的模块顺序） ──
        optimized = optimized_data.get("optimized", {})

        # 第一步：按原始简历的模块顺序排列（Python dict 保持插入顺序）
        all_sections = []
        seen = set()
        for s in original_resume.keys():
            if s not in seen and s not in ("name", "email", "phone", "基本信息"):
                if isinstance(original_resume.get(s), str):
                    all_sections.append(s)
                    seen.add(s)

        # 第二步：补充原始简历中没有、但优化后新增的模块
        for s in optimized.keys():
            if s not in seen:
                all_sections.append(s)
                seen.add(s)

        for section in all_sections:
            content = None
            is_optimized = False
            if section in optimized:
                content = _clean_llm_output(optimized[section])
                is_optimized = True
            elif section in original_resume and isinstance(original_resume[section], str):
                content = original_resume[section]

            if content:
                self._add_professional_section(doc, section, content, is_optimized)

        doc.save(output_path)
        return str(output_path)

    def generate_markdown(self, optimized_data: Dict, original_resume: Dict) -> str:
        """生成 Markdown 格式的简历（保持原始简历模块顺序）"""
        lines = ["# 个人简历", ""]

        if original_resume.get("name"):
            lines.append(f"**{original_resume['name']}**")
        contact = []
        if original_resume.get("email"):
            contact.append(original_resume['email'])
        if original_resume.get("phone"):
            contact.append(original_resume['phone'])
        if contact:
            lines.append(" | ".join(contact))
        lines.append("")

        optimized = optimized_data.get("optimized", {})

        # 按原始简历顺序排列模块
        all_sections = []
        seen = set()
        for s in original_resume.keys():
            if s not in seen and s not in ("name", "email", "phone", "基本信息"):
                all_sections.append(s)
                seen.add(s)
        for s in optimized.keys():
            if s not in seen:
                all_sections.append(s)
                seen.add(s)

        for section in all_sections:
            content = optimized.get(section) or original_resume.get(section)
            if content and isinstance(content, str):
                content = _clean_llm_output(content)
                lines.append(f"## {section}")
                lines.append(content)
                lines.append("")

        return "\n".join(lines)

    def generate_report(self, optimized_data: Dict, match_result=None, jobs: List = None) -> str:
        """生成优化报告（完整展示，不截断）"""
        lines = [
            "# 简历优化报告",
            "",
            f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
        ]

        # 匹配信息
        if match_result:
            lines.extend([
                "## 职位匹配分析",
                "",
                f"| 维度 | 结果 |",
                f"|------|------|",
                f"| 目标职位 | {match_result.job_title} @ {match_result.company} |",
                f"| 综合匹配度 | {match_result.overall_score * 100:.1f}% |",
                f"| 技能匹配度 | {match_result.skill_match_score * 100:.1f}% |",
                f"| 经验匹配度 | {match_result.experience_match_score * 100:.1f}% |",
                "",
            ])

            if match_result.matched_skills:
                lines.append(f"**已匹配技能**：{', '.join(match_result.matched_skills)}")
                lines.append("")

            if match_result.missing_skills:
                lines.append(f"**缺少技能**：{', '.join(match_result.missing_skills)}")
                lines.append("")

            if match_result.suggestions:
                lines.extend(["### 优化建议", ""])
                for s in match_result.suggestions:
                    lines.append(f"- {s}")
                lines.append("")

        # 优化说明
        notes = optimized_data.get("optimization_notes", [])
        if notes:
            lines.extend(["## 优化说明", ""])
            for note in notes:
                lines.append(f"- {note}")
            lines.append("")

        # 警告
        report = optimized_data.get("report", {})
        warnings = report.get("warnings", [])
        if warnings:
            lines.extend(["## 注意事项", ""])
            for w in warnings:
                lines.append(f"> {w}")
            lines.append("")

        # 优化内容对比（完整展示）
        lines.extend(["## 优化内容对比", ""])
        optimized = optimized_data.get("optimized", {})
        original = optimized_data.get("original", {})

        for section, content in optimized.items():
            content = _clean_llm_output(content)
            orig_content = _clean_llm_output(original.get(section, ""))

            lines.append(f"### {section}")
            lines.append("")

            if orig_content:
                lines.append("<details>")
                lines.append(f"<summary>原文 ({len(orig_content)}字)</summary>")
                lines.append("")
                lines.append(orig_content)
                lines.append("")
                lines.append("</details>")
                lines.append("")

            lines.append(f"**优化后** ({len(content)}字)：")
            lines.append("")
            lines.append(content)
            lines.append("")

        # 参考职位信息
        if jobs:
            lines.extend(["## 参考职位信息", ""])
            lines.append("| # | 职位名称 | 公司名称 | 地点 | 链接 |")
            lines.append("|---|---------|---------|------|------|")
            for i, job in enumerate(jobs, 1):
                title = job.title or "未知"
                company = job.company or "未知"
                location = job.location or "-"
                url = job.url or "-"
                if url and url != "-":
                    lines.append(f"| {i} | {title} | {company} | {location} | [查看]({url}) |")
                else:
                    lines.append(f"| {i} | {title} | {company} | {location} | - |")
            lines.append("")

        return "\n".join(lines)

    # ── DOCX 排版辅助方法 ──────────────────────────────────

    def _setup_page(self, doc: Document):
        """设置页面：A4、合理页边距"""
        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    def _setup_styles(self, doc: Document):
        """设置默认样式"""
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10)
        style.font.color.rgb = self.SECONDARY_COLOR
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.25
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _add_name_header(self, doc: Document, name: str):
        """添加姓名大标题"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(name)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _add_contact_row(self, doc: Document, resume_data: Dict):
        """添加联系方式（一行居中，竖线分隔）"""
        parts = []

        # 从基本信息模块中提取性别、年龄、工作经验等（排除姓名、电话、邮箱）
        basic_info = resume_data.get("基本信息", "")
        if basic_info:
            name = self._extract_name(resume_data)
            phone = resume_data.get("phone", "")
            email = resume_data.get("email", "")
            # 用正则拆分所有 token，按分隔符切分
            tokens = re.split(r'[\s|｜]+', basic_info)
            # 过滤掉姓名、电话、邮箱，保留其他信息（性别、年龄、工作经验等）
            skip = {name, phone, email}
            for t in tokens:
                t = t.strip()
                if t and t not in skip:
                    parts.append(t)

        if resume_data.get("phone"):
            parts.append(resume_data["phone"])
        if resume_data.get("email"):
            parts.append(resume_data["email"])

        if parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            text = "  |  ".join(parts)
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = self.LIGHT_GRAY
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 分隔线
        self._add_divider(doc)

    def _add_divider(self, doc: Document):
        """添加水平分隔线"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        # 使用底部边框模拟分隔线
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{self.SECTION_BG}"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    @staticmethod
    def _strip_inline_markdown(text: str) -> str:
        """去除行内的 Markdown 加粗/斜体标记，只保留纯文本"""
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # **bold**
        text = re.sub(r'\*(.+?)\*', r'\1', text)       # *italic*
        return text

    def _add_professional_section(self, doc: Document, title: str, content: str, is_optimized: bool):
        """添加专业排版的简历模块"""
        # 模块标题（带左侧色块）
        self._add_section_heading(doc, title)

        # 教育经历使用专用渲染（自动拆分多条经历）
        section_key = title.replace(" ", "")
        if any(kw in section_key for kw in ("教育", "学历")):
            self._render_education_section(doc, content)
            return

        # 解析并添加内容
        lines = content.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # 先去除 Markdown 标记再检测行类型
            clean = self._strip_inline_markdown(stripped)

            # 检测是否是子标题（如公司名称+职位+时间 这种行）
            if self._is_sub_heading(clean):
                self._add_sub_heading(doc, clean)
                continue

            # 检测是否是列表项
            if self._is_list_item(clean):
                text = self._clean_list_prefix(clean)
                self._add_bullet_item(doc, text)
                continue

            # 普通段落
            self._add_body_text(doc, clean)

        # 模块间间距
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def _add_section_heading(self, doc: Document, title: str):
        """带左侧色块的模块标题"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)

        # 添加左侧色块标记
        run = p.add_run("■ ")
        run.font.size = Pt(10)
        run.font.color.rgb = self.PRIMARY_COLOR

        run = p.add_run(title)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = self.PRIMARY_COLOR
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 标题下方细线
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="D0D0D0"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _add_sub_heading(self, doc: Document, text: str):
        """添加子标题（公司/学校名 + 职位 + 时间）"""
        # 先清理残留的 Markdown 标记
        text = self._strip_inline_markdown(text).strip()
        # 清理可能的残留 * 符号
        text = re.sub(r'(?<!\*)\*(?!\*)', '', text)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

        # 尝试拆分：公司名 | 职位 | 时间
        parts = re.split(r'\s*[|｜]\s*|\s{2,}', text)
        if len(parts) >= 2:
            # 第一部分加粗加大（公司/学校名）
            run = p.add_run(parts[0].strip())
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = self.SECONDARY_COLOR
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

            # 其余部分正常显示
            for part in parts[1:]:
                part = part.strip()
                if not part:
                    continue
                run = p.add_run(f"  {part}")
                run.font.size = Pt(10)
                run.font.color.rgb = self.LIGHT_GRAY
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        else:
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = self.SECONDARY_COLOR
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _render_education_section(self, doc: Document, content: str):
        """教育经历专用渲染：自动拆分多条经历为独立子标题+详情"""
        text = self._strip_inline_markdown(content).strip()

        # 按学位关键词拆分多条教育经历（token 扫描法）
        tokens = text.split()
        degree_keywords = {'硕士', '本科', '博士', '学士'}
        split_indices = []
        seen_date = False  # 是否已遇到过日期范围

        for i in range(len(tokens)):
            # 检测日期范围 token（如 2022-2024）
            if re.match(r'\d{4}[\-./–]\d{1,4}', tokens[i]):
                seen_date = True
            # 学位关键词 + 下一个 token 是中文学校名 → 新条目起点
            # 但必须在已见过日期之后（避免把同一条目内的"学校 硕士 专业"拆开）
            if tokens[i] in degree_keywords and seen_date and i + 1 < len(tokens):
                next_tok = tokens[i + 1]
                if re.match(r'^[\u4e00-\u9fff]{2,}$', next_tok):
                    split_indices.append(i)

        # 按 split_indices 拆分 token 列表
        raw_entries = []
        prev = 0
        for idx in split_indices:
            entry_tokens = tokens[prev:idx]
            # 检查尾部是否包含下一条的学校名（在日期之后出现的中文 token）
            # 例如: "...2022-2024 两年制，绩点3.95/4 西南交通大学"
            #       → "西南交通大学" 应属于下一条
            tail_school = None
            for j in range(len(entry_tokens) - 1, -1, -1):
                if re.match(r'\d{4}[\-./–]\d{1,4}', entry_tokens[j]):
                    # 日期之后的中文 token 可能是下一条的学校名
                    after_date = entry_tokens[j + 1:]
                    chinese_after = [t for t in after_date if re.match(r'^[\u4e00-\u9fff]{2,}$', t)]
                    if chinese_after:
                        tail_school = chinese_after[-1]  # 取最后一个中文 token
                    break
            if tail_school:
                # 从当前条目中移除学校名及其之后的内容
                school_pos = None
                for j in range(len(entry_tokens) - 1, -1, -1):
                    if entry_tokens[j] == tail_school:
                        school_pos = j
                        break
                if school_pos is not None:
                    entry_tokens = entry_tokens[:school_pos]
                    # 把学校名加到下一条的开头
                    tokens[idx:idx] = [tail_school]
                    # 调整后续 split_indices
                    split_indices = [si + 1 if si > idx else si for si in split_indices]

            raw_entries.append(' '.join(entry_tokens))
            prev = idx
        raw_entries.append(' '.join(tokens[prev:]))
        raw_entries = [e.strip() for e in raw_entries if e.strip()]

        # 如果只拆出1条，尝试按日期范围拆
        if len(raw_entries) <= 1:
            date_matches = list(re.finditer(r'\d{4}[\-./–]\d{1,4}(?:\s*至今)?', text))
            if len(date_matches) >= 2:
                # 有多个日期范围但没被学位关键词拆分，回退到简单日期拆分
                raw_entries = []
                for i, dm in enumerate(date_matches):
                    start = date_matches[i - 1].end() if i > 0 else 0
                    end = dm.end()
                    raw_entries.append(text[start:end].strip())

        if len(raw_entries) <= 1:
            # 确实只有1条经历，尝试提取子标题
            date_match = re.search(r'\d{4}[\-./–]\d{1,4}(?:\s*至今)?', text)
            if date_match:
                self._add_sub_heading(doc, text[:date_match.end()].strip())
                detail = text[date_match.end():].strip()
                if detail:
                    self._add_body_text(doc, detail)
            else:
                self._add_body_text(doc, text)
            self._add_section_spacing(doc)
            return

        # 多条经历：逐条解析并渲染
        for entry_text in raw_entries:
            # 用正则提取：学校名 学位 专业 日期 [详情]
            m = re.match(
                r'([\u4e00-\u9fff\w]+(?:[\u4e00-\u9fff\w（）()]*?))\s+'
                r'(硕士|本科|博士|学士)\s+'
                r'(.+?)\s+'
                r'(\d{4}[\-./–]\d{1,4}(?:\s*至今)?)'
                r'(.*)',
                entry_text
            )
            if m:
                school, degree, major = m.group(1), m.group(2), m.group(3)
                date_str, detail = m.group(4), m.group(5).strip()
                self._add_sub_heading(doc, f"{school}  {degree} {major}  {date_str}")
                if detail:
                    self._add_body_text(doc, detail)
            else:
                # 回退：直接用日期范围作为子标题
                date_match = re.search(r'\d{4}[\-./–]\d{1,4}(?:\s*至今)?', entry_text)
                if date_match:
                    self._add_sub_heading(doc, entry_text[:date_match.end()].strip())
                    detail = entry_text[date_match.end():].strip()
                    if detail:
                        self._add_body_text(doc, detail)
                else:
                    self._add_body_text(doc, entry_text)

        self._add_section_spacing(doc)

    def _add_section_spacing(self, doc: Document):
        """添加模块间间距"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    def _add_bullet_item(self, doc: Document, text: str):
        """添加带圆点的列表项"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.2

        # 圆点标记
        run = p.add_run("• ")
        run.font.size = Pt(10)
        run.font.color.rgb = self.ACCENT_COLOR
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 处理加粗标记 **text**
        self._add_rich_text(doc, p, text)

    def _add_body_text(self, doc: Document, text: str):
        """添加正文段落"""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.2
        self._add_rich_text(doc, p, text)

    def _add_rich_text(self, doc: Document, paragraph, text: str):
        """处理文本中的 Markdown 加粗标记 **text**"""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            else:
                run = paragraph.add_run(part)
            run.font.size = Pt(10)
            run.font.color.rgb = self.SECONDARY_COLOR
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    @staticmethod
    def _extract_name(resume_data: Dict) -> str:
        """
        从简历数据中提取姓名。
        优先使用 name 字段，否则从基本信息模块中提取第一行第一个词。
        """
        name = resume_data.get("name")
        if name:
            return name.strip()
        basic_info = resume_data.get("基本信息", "")
        if basic_info:
            # 取第一行，再取第一个非空词（姓名通常在第一个位置）
            first_line = basic_info.split('\n')[0].strip()
            # 按常见分隔符拆分
            parts = re.split(r'[\s|｜,，]+', first_line)
            if parts:
                candidate = parts[0].strip()
                # 姓名通常是2-4个中文字符
                if candidate and re.match(r'^[\u4e00-\u9fff]{2,4}$', candidate):
                    return candidate
        return ""

    @staticmethod
    def _is_sub_heading(line: str) -> bool:
        """判断是否是子标题行（包含公司/学校+时间等信息）"""
        # 包含日期格式的行（如 2025.02-2025.11、2025/03-2025/11）
        if re.search(r'\d{4}[./\-]\d{1,2}', line):
            # 但不能是纯列表项
            if not line.startswith(('-', '*', '•')) and not line[0].isdigit():
                return True
            # 以数字开头但后面跟着公司名（如 "1. 公司名 2025.02-..."）
            if re.match(r'^\d+\.\s*\S+.*\d{4}[./\-]\d{1,2}', line):
                return True
        return False

    @staticmethod
    def _is_list_item(line: str) -> bool:
        """判断是否是列表项"""
        if line.startswith(('-', '*', '•')):
            return True
        # 数字编号：1. 2. 3. 或 ① ② ③
        if re.match(r'^[\d]+[.、]\s', line):
            return True
        if re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', line):
            return True
        return False

    @staticmethod
    def _clean_list_prefix(line: str) -> str:
        """去除列表前缀"""
        line = re.sub(r'^[-*•]\s*', '', line)
        line = re.sub(r'^[\d]+[.、]\s*', '', line)
        line = re.sub(r'^[①②③④⑤⑥⑦⑧⑨⑩]\s*', '', line)
        return line.strip()
