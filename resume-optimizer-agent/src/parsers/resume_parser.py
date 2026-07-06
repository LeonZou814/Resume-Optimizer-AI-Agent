"""
简历解析模块
支持 PDF、DOCX、TXT 格式
"""

import io
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

import pdfplumber
from docx import Document

from src.utils.text_utils import clean_text, split_by_sections


@dataclass
class ResumeData:
    """简历数据结构"""
    raw_text: str = ""
    sections: dict = field(default_factory=dict)
    file_type: str = ""
    file_path: Optional[str] = None
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None


class ResumeParser:
    """简历解析器"""

    def parse(self, file_path: str | Path) -> ResumeData:
        """根据文件类型自动选择解析方法"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix == ".docx":
            return self._parse_docx(path)
        elif suffix == ".txt":
            return self._parse_txt(path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    def parse_bytes(self, file_bytes: bytes, filename: str) -> ResumeData:
        """从字节流解析简历"""
        suffix = Path(filename).suffix.lower()
        file_type = suffix.lstrip(".")

        if suffix == ".pdf":
            text = self._extract_pdf_text(io.BytesIO(file_bytes))
        elif suffix == ".docx":
            text = self._extract_docx_text(io.BytesIO(file_bytes))
        elif suffix == ".txt":
            text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

        return self._build_resume_data(text, file_type, filename)

    def _parse_pdf(self, path: Path) -> ResumeData:
        """解析 PDF 简历"""
        text = self._extract_pdf_text(str(path))
        return self._build_resume_data(text, "pdf", str(path))

    def _extract_pdf_text(self, source) -> str:
        """从 PDF 提取文本"""
        texts = []
        with pdfplumber.open(source) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    texts.append(page_text)
        return "\n".join(texts)

    def _parse_docx(self, path: Path) -> ResumeData:
        """解析 DOCX 简历"""
        text = self._extract_docx_text(str(path))
        return self._build_resume_data(text, "docx", str(path))

    def _extract_docx_text(self, source) -> str:
        """从 DOCX 提取文本"""
        doc = Document(source)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也读取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    def _parse_txt(self, path: Path) -> ResumeData:
        """解析 TXT 简历"""
        text = path.read_text(encoding="utf-8", errors="ignore")
        return self._build_resume_data(text, "txt", str(path))

    def _build_resume_data(self, raw_text: str, file_type: str, file_path: str) -> ResumeData:
        """构建简历数据对象"""
        # 先拆分模块（需要保留换行符），再清理每个模块的内容
        sections = split_by_sections(raw_text)

        # 清理原始文本（用于向量匹配等）
        clean = clean_text(raw_text)

        resume = ResumeData(
            raw_text=clean,
            sections=sections,
            file_type=file_type,
            file_path=file_path,
        )

        # 尝试提取联系方式
        resume.email = self._extract_email(clean)
        resume.phone = self._extract_phone(clean)

        return resume

    @staticmethod
    def _extract_email(text: str) -> Optional[str]:
        import re
        match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        return match.group(0) if match else None

    @staticmethod
    def _extract_phone(text: str) -> Optional[str]:
        import re
        match = re.search(r'1[3-9]\d{9}', text)
        return match.group(0) if match else None
